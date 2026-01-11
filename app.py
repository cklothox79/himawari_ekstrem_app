# =========================================================
# 🛰️ HIMAWARI EXTREME WEATHER ANALYSIS – STAGE 3
# Time Series + PDF & DOCX Export
# =========================================================

import streamlit as st
import xarray as xr
import numpy as np
import pandas as pd
from pathlib import Path
import matplotlib.pyplot as plt
import re
from io import BytesIO

# PDF
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table
from reportlab.lib.styles import getSampleStyleSheet

# DOCX
from docx import Document

# =========================================================
# KONFIGURASI
# =========================================================
st.set_page_config(page_title="Analisis TBB Himawari", layout="centered")

BASE_DIR = Path(__file__).parent
DATA_DIR = BASE_DIR / "data_nc"

# =========================================================
# FUNGSI
# =========================================================
def find_nearest_pixel(lat2d, lon2d, lat0, lon0):
    dist2 = (lat2d - lat0)**2 + (lon2d - lon0)**2
    iy, ix = np.unravel_index(np.argmin(dist2), dist2.shape)
    return iy, ix


def extract_mean_tbb(ds, lat0, lon0, radius_km):
    tbb = ds["tbb"].values
    lat = ds["latitude"].values
    lon = ds["longitude"].values

    if lat.ndim == 1:
        lon2d, lat2d = np.meshgrid(lon, lat)
    else:
        lat2d, lon2d = lat, lon

    iy, ix = find_nearest_pixel(lat2d, lon2d, lat0, lon0)
    radius_px = int(radius_km / 2)

    y0, y1 = max(iy - radius_px, 0), min(iy + radius_px, tbb.shape[0])
    x0, x1 = max(ix - radius_px, 0), min(ix + radius_px, tbb.shape[1])

    return np.nanmean(tbb[y0:y1, x0:x1]) - 273.15


def parse_band_time(fname):
    band = re.search(r"B(\d{2})", fname).group(1)
    time = re.search(r"_(\d{12})\.nc", fname).group(1)
    return band, time


def interpret_tbb(min_tbb):
    if min_tbb < -60:
        return "Awan Cumulonimbus sangat intens dengan potensi cuaca ekstrem tinggi."
    elif min_tbb < -50:
        return "Awan Cumulonimbus matang, berpotensi menimbulkan cuaca ekstrem."
    elif min_tbb < -40:
        return "Awan konvektif berkembang."
    else:
        return "Awan non-konvektif dominan."


# =========================================================
# UI
# =========================================================
st.title("🛰️ Analisis Suhu Puncak Awan (TBB) – Himawari")
st.caption("Tahap 3 – Time Series + Export Laporan")

lat0 = st.number_input("Lintang (°)", value=-7.3735, format="%.4f")
lon0 = st.number_input("Bujur (°)", value=112.7938, format="%.4f")
radius_km = st.slider("Radius Analisis (km)", 2, 20, 10)

# =========================================================
# PROSES
# =========================================================
if st.button("▶️ Jalankan Analisis"):
    records = []

    for f in sorted(DATA_DIR.glob("*.nc")):
        try:
            band, time = parse_band_time(f.name)
            ds = xr.open_dataset(f)
            mean_tbb = extract_mean_tbb(ds, lat0, lon0, radius_km)

            records.append({
                "Waktu (UTC)": time,
                "Band": f"B{band}",
                "Mean TBB (°C)": round(mean_tbb, 2)
            })
        except:
            pass

    df = pd.DataFrame(records)
    df["Waktu (UTC)"] = pd.to_datetime(df["Waktu (UTC)"], format="%Y%m%d%H%M")
    df = df.sort_values("Waktu (UTC)")

    st.subheader("📊 Tabel Time Series")
    st.dataframe(df, use_container_width=True)

    # Grafik
    fig, ax = plt.subplots()
    for band in df["Band"].unique():
        d = df[df["Band"] == band]
        ax.plot(d["Waktu (UTC)"], d["Mean TBB (°C)"], marker="o", label=band)
    ax.legend()
    ax.grid(True)
    st.pyplot(fig)

    # Interpretasi
    min_tbb = df["Mean TBB (°C)"].min()
    interpretation = interpret_tbb(min_tbb)

    st.subheader("📝 Interpretasi")
    st.write(interpretation)

    # =====================================================
    # EXPORT
    # =====================================================
    st.subheader("📄 Unduh Laporan")

    # ---- PDF ----
    pdf_buffer = BytesIO()
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(pdf_buffer)
    story = [
        Paragraph("<b>LAPORAN ANALISIS CUACA EKSTREM</b>", styles["Title"]),
        Spacer(1, 12),
        Paragraph(f"Lokasi: {lat0}, {lon0}", styles["Normal"]),
        Paragraph(f"Radius Analisis: {radius_km} km", styles["Normal"]),
        Spacer(1, 12),
        Paragraph("<b>Interpretasi:</b>", styles["Heading2"]),
        Paragraph(interpretation, styles["Normal"])
    ]
    table_data = [df.columns.tolist()] + df.values.tolist()
    story.append(Spacer(1, 12))
    story.append(Table(table_data))
    doc.build(story)

    st.download_button(
        "⬇️ Download PDF",
        data=pdf_buffer.getvalue(),
        file_name="laporan_tbb_himawari.pdf",
        mime="application/pdf"
    )

    # ---- DOCX ----
    docx = Document()
    docx.add_heading("LAPORAN ANALISIS CUACA EKSTREM", level=1)
    docx.add_paragraph(f"Lokasi: {lat0}, {lon0}")
    docx.add_paragraph(f"Radius Analisis: {radius_km} km")
    docx.add_heading("Interpretasi", level=2)
    docx.add_paragraph(interpretation)

    table = docx.add_table(rows=1, cols=len(df.columns))
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = col

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    buffer = BytesIO()
    docx.save(buffer)

    st.download_button(
        "⬇️ Download DOCX",
        data=buffer.getvalue(),
        file_name="laporan_tbb_himawari.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
